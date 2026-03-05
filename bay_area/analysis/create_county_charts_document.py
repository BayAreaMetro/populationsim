#!/usr/bin/env python3
"""
Create Word Document with County Analysis Charts
==============================================

Combines all county analysis charts into a single Word document
formatted for landscape 8.5x11 pages with proper sizing and organization.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd

def create_county_charts_document():
    """Create Word document with all county analysis charts"""
    
    print("="*80)
    print("CREATING COUNTY ANALYSIS CHARTS DOCUMENT")
    print("="*80)
    
    # Paths
    charts_dir = Path("output_2023/charts/county_analysis")
    output_file = Path("output_2023/County_Analysis_Charts_Report.docx")
    
    if not charts_dir.exists():
        print(f"Error: Charts directory not found: {charts_dir}")
        return
    
    # Get all PNG files
    png_files = list(charts_dir.glob("*.png"))
    png_files.sort()
    
    if not png_files:
        print("No PNG files found in charts directory")
        return
    
    print(f"Found {len(png_files)} chart files")
    
    # Create Word document
    doc = Document()
    
    # Set up landscape orientation for all sections
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
    
    # Add title page
    title = doc.add_heading("PopulationSim 2023 County Analysis", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph("County-Level Performance and Distribution Analysis")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0].font
    subtitle_format.size = Pt(16)
    subtitle_format.bold = True
    
    date_para = doc.add_paragraph("Generated: November 19, 2025")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("\n" * 3)
    
    overview_text = doc.add_paragraph(
        "This report contains county-level performance analysis for the Bay Area 9-county region. "
        "Charts show control versus result comparisons, performance metrics, and demographic distributions "
        "at the county geography level."
    )
    overview_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Add summary tables if available
    summary_file = charts_dir / "county_performance_summary.csv"
    detailed_file = charts_dir / "county_detailed_results.csv"
    
    if summary_file.exists():
        doc.add_page_break()
        
        # Add performance summary table
        summary_heading = doc.add_heading("County Performance Summary", level=1)
        summary_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        try:
            df = pd.read_csv(summary_file)
            
            # Add table with key columns
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'County'
            hdr_cells[1].text = 'Control Total'
            hdr_cells[2].text = 'Result Total'
            hdr_cells[3].text = 'Difference'
            hdr_cells[4].text = 'Diff %'
            hdr_cells[5].text = 'MAE'
            
            # Data rows
            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(row.get('county_name', ''))
                row_cells[1].text = f"{row.get('total_control', 0):,.0f}"
                row_cells[2].text = f"{row.get('total_result', 0):,.0f}"
                row_cells[3].text = f"{row.get('total_diff', 0):,.0f}"
                row_cells[4].text = f"{row.get('total_diff_pct', 0):.3f}%"
                row_cells[5].text = f"{row.get('mae', 0):,.1f}"
            
            print(f"  ✓ Added county performance summary table")
            
        except Exception as e:
            print(f"  ⚠ Could not add summary table: {e}")
    
    # Add page break before charts
    doc.add_page_break()
    
    # Organize charts by category
    chart_categories = {
        "Performance Overview": [
            "county_performance_overview.png",
            "county_variable_heatmap.png"
        ],
        "Demographic Analysis": [
            "county_age_analysis.png",
            "county_household_size_analysis.png",
            "county_workers_analysis.png",
            "county_income_analysis.png"
        ],
        "Detailed Variables": [f for f in png_files if f.name not in [
            "county_performance_overview.png",
            "county_variable_heatmap.png", 
            "county_age_analysis.png",
            "county_household_size_analysis.png",
            "county_workers_analysis.png",
            "county_income_analysis.png"
        ]]
    }
    
    for category, chart_patterns in chart_categories.items():
        # Find matching files
        category_files = []
        for pattern in chart_patterns:
            if isinstance(pattern, str):
                # String pattern - find exact match
                matching_files = [f for f in png_files if f.name == pattern]
            else:
                # Path object - direct match
                matching_files = [pattern] if pattern in png_files else []
            category_files.extend(matching_files)
        
        if not category_files:
            continue
            
        # Add category header
        category_heading = doc.add_heading(category, level=1)
        category_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add charts for this category
        charts_added = 0
        for chart_file in category_files:
            if not chart_file.exists():
                continue
                
            # Add chart title (derived from filename)
            chart_name = chart_file.stem.replace("county_", "").replace("_analysis", "").replace("_", " ").title()
            chart_title = doc.add_heading(chart_name, level=2)
            chart_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add image - size to fit landscape page with margins
            # Landscape usable area: ~10 inches wide, ~7 inches tall
            try:
                doc.add_picture(str(chart_file), width=Inches(9.5))
                charts_added += 1
                print(f"  ✓ Added chart: {chart_file.name}")
                
                # Add page break after each chart except the last in category
                if charts_added < len(category_files):
                    doc.add_page_break()
                    
            except Exception as e:
                print(f"  ⚠ Could not add chart {chart_file.name}: {e}")
                continue
        
        # Add page break between categories
        if category != list(chart_categories.keys())[-1]:
            doc.add_page_break()
    
    # Save document
    doc.save(output_file)
    
    print(f"\n{'='*80}")
    print(f"✅ County Analysis Charts Document created: {output_file}")
    print(f"📊 Document contains {len(png_files)} charts organized by category")
    print(f"📄 Formatted for landscape 8.5x11 pages")
    print(f"{'='*80}")

if __name__ == '__main__':
    create_county_charts_document()
