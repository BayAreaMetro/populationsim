#!/usr/bin/env python3
"""
Create Word Document with TAZ Analysis Charts
============================================

Combines all TAZ analysis charts into a single Word document
formatted for landscape 8.5x11 pages with proper sizing and organization.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd

def create_taz_charts_document():
    """Create Word document with all TAZ analysis charts"""
    
    print("="*80)
    print("CREATING TAZ ANALYSIS CHARTS DOCUMENT")
    print("="*80)
    
    # Paths
    charts_dir = Path("output_2023/charts/taz_analysis")
    output_file = Path("output_2023/TAZ_Analysis_Charts_Report.docx")
    
    if not charts_dir.exists():
        print(f"Error: Charts directory not found: {charts_dir}")
        return
    
    # Get all PNG files
    png_files = list(charts_dir.glob("*.png"))
    png_files.sort()
    
    if not png_files:
        print("No PNG files found in charts directory")
        return
    
    print(f"Found {len(png_files)} chart files")
    
    # Create Word document
    doc = Document()
    
    # Set up landscape orientation for all sections
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
    
    # Add title page
    title = doc.add_heading("PopulationSim 2023 TAZ Analysis", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph("Controls vs Results Distribution Analysis")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0].font
    subtitle_format.size = Pt(16)
    subtitle_format.bold = True
    
    date_para = doc.add_paragraph("Generated: November 3, 2025")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("\\n" * 3)
    
    overview_text = doc.add_paragraph(
        "This report contains scatter plot analyses comparing control targets "
        "versus PopulationSim results for all variables at the Traffic Analysis Zone (TAZ) level. "
        "Each chart shows the relationship between input controls and synthesized outputs, "
        "with performance metrics including R-squared values and best-fit line equations."
    )
    overview_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Add page break
    doc.add_page_break()
    
    # Organize charts by category
    chart_categories = {
        "Population and Households": ["taz_numhh_analysis.png", "taz_numhh_gq_analysis.png", 
                                    "taz_total_population_analysis.png", "taz_overall_summary.png"],
        "Household Size": [f for f in png_files if "hh_size_" in f.name],
        "Workers": [f for f in png_files if "hh_wrks_" in f.name],
        "Age Groups": [f for f in png_files if "pers_age_" in f.name],
        "Income": [f for f in png_files if "inc_" in f.name],
        "Group Quarters": [f for f in png_files if "hh_gq_" in f.name and "analysis" in f.name],
        "Household Composition": [f for f in png_files if "hh_kids_" in f.name]
    }
    
    for category, chart_patterns in chart_categories.items():
        # Find matching files
        category_files = []
        for pattern in chart_patterns:
            if isinstance(pattern, str):
                # String pattern - find exact match
                matching_files = [f for f in png_files if f.name == pattern]
            else:
                # Path object - direct match
                matching_files = [pattern] if pattern in png_files else []
            category_files.extend(matching_files)
        
        if not category_files:
            continue
            
        # Add category header
        category_heading = doc.add_heading(category, level=1)
        category_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add charts for this category
        charts_added = 0
        for chart_file in category_files:
            if not chart_file.exists():
                continue
                
            # Add chart title (derived from filename)
            chart_name = chart_file.stem.replace("taz_", "").replace("_analysis", "").replace("_", " ").title()
            chart_title = doc.add_heading(chart_name, level=2)
            chart_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add image - size to fit landscape page with margins
            # Landscape usable area: ~10 inches wide, ~7 inches tall
            try:
                doc.add_picture(str(chart_file), width=Inches(9.5))
                charts_added += 1
                
                # Add page break after each chart except the last in category
                if charts_added < len(category_files):
                    doc.add_page_break()
                    
            except Exception as e:
                print(f"Warning: Could not add chart {chart_file.name}: {e}")
                continue
        
        # Add page break between categories
        if category != list(chart_categories.keys())[-1]:
            doc.add_page_break()
    
    # Add summary statistics if available
    summary_file = charts_dir / "taz_analysis_summary.csv"
    if summary_file.exists():
        doc.add_page_break()
        summary_heading = doc.add_heading("Performance Summary", level=1)
        summary_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        try:
            df = pd.read_csv(summary_file)
            
            # Add summary table
            table = doc.add_table(rows=1, cols=len(df.columns))
            table.style = 'Table Grid'
            
            # Header row
            header_cells = table.rows[0].cells
            for i, column in enumerate(df.columns):
                header_cells[i].text = str(column)
                header_cells[i].paragraphs[0].runs[0].font.bold = True
            
            # Data rows
            for _, row in df.head(20).iterrows():  # Limit to first 20 rows
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)
                    
        except Exception as e:
            doc.add_paragraph(f"Error loading summary table: {e}")
    
    # Save document
    doc.save(output_file)
    
    print(f"✅ TAZ Analysis Charts Document created: {output_file}")
    print(f"📊 Document contains {len(png_files)} charts organized by category")
    print(f"📄 Formatted for landscape 8.5x11 pages")
    
    return output_file

if __name__ == '__main__':
    create_taz_charts_document()